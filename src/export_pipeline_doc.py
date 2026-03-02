"""
Generates a formatted .docx document describing the complete
SoilScan Sentinel-2 data pipeline from the researcher's perspective.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUTPUT_PATH = "outputs/SoilScan_Pipeline_Documentation.docx"

# ── Helpers ──────────────────────────────────────────────────────────────────

def set_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    run = h.runs[0]
    if level == 1:
        run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
        run.font.size = Pt(15)
    elif level == 2:
        run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
        run.font.size = Pt(12)
    else:
        run.font.color.rgb = RGBColor(0x40, 0x40, 0x40)
        run.font.size = Pt(11)
    return h

def add_para(doc, text, bold=False, italic=False, size=Pt(11), indent=False):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = size
    return p

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.3 + level * 0.25)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    return p

def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        run = hdr_cells[i].paragraphs[0].runs[0]
        run.bold = True
        run.font.size = Pt(10)
        tc = hdr_cells[i]._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "2E74B5")
        tcPr.append(shd)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for r, row_data in enumerate(rows):
        cells = table.rows[r + 1].cells
        for c, val in enumerate(row_data):
            cells[c].text = str(val)
            cells[c].paragraphs[0].runs[0].font.size = Pt(10)
    return table

def add_code(doc, code_text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run(code_text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1E, 0x1E, 0x1E)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F0F0F0")
    pPr.append(shd)
    return p

def add_divider(doc):
    doc.add_paragraph()


# ── Build Document ────────────────────────────────────────────────────────────

doc = Document()

for section in doc.sections:
    section.top_margin    = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin   = Inches(1.2)
    section.right_margin  = Inches(1.2)

# ── Title ─────────────────────────────────────────────────────────────────────
title = doc.add_heading("SoilScan Sentinel-2", 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

sub = doc.add_paragraph("Complete Data Collection and Processing Pipeline")
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.runs[0].font.size = Pt(13)
sub.runs[0].italic = True

doc.add_paragraph()

add_para(doc,
    "This document describes the complete methodology followed by the researchers "
    "in the SoilScan project. The pipeline covers physical field data collection "
    "in Benguet, Philippines, satellite imagery acquisition and processing via "
    "the Copernicus Data Space Ecosystem, and the training of machine learning models "
    "for soil nutrient prediction using Sentinel-2 spectral data.")

add_divider(doc)

# ── Phase 1 ───────────────────────────────────────────────────────────────────
set_heading(doc, "Phase 1 — Field Data Collection", level=1)

add_para(doc,
    "The researchers conducted field surveys across multiple farm sites in Benguet, "
    "Philippines. At each survey spot, data was recorded using the AgriCapture mobile "
    "application, which was developed to standardize the collection of both spatial "
    "and environmental parameters.")

set_heading(doc, "Data Recorded per Field Spot", level=2)
add_bullet(doc, "GPS coordinates — latitude, longitude, and altitude (meters above sea level)")
add_bullet(doc, "Photograph — a top-down image of the field and soil surface")
add_bullet(doc, "Microclimate readings — air temperature (degrees Celsius) and relative humidity (%)")
add_bullet(doc, "Field metadata — municipality, barangay, farm name, and crop type")
add_bullet(doc, "Capture timestamp — exact date and time in UTC")
add_bullet(doc, "Device metadata — camera pitch, roll, and heading; GPS accuracy; device ID")

set_heading(doc, "Output", level=2)
add_para(doc,
    "All field observations were compiled and exported as a structured CSV file "
    "containing one row per captured field spot:")
add_code(doc, "data/external/combined_field_data.csv")
add_para(doc,
    "The dataset consists of 1,254 rows collected across four barangays: "
    "Balili, Paoay, and Alapang in La Trinidad, and Betag in Atok, Benguet.",
    indent=True)

add_divider(doc)

# ── Phase 2 ───────────────────────────────────────────────────────────────────
set_heading(doc, "Phase 2 — Sentinel-2 Satellite Data Acquisition and Band Extraction", level=1)

add_para(doc,
    "To obtain spectral reflectance data corresponding to each field spot, the researchers "
    "developed an automated pipeline (data_fetcher_copernicus.py) that searches for, "
    "downloads, and samples Sentinel-2 L2A satellite imagery from the "
    "Copernicus Data Space Ecosystem (CDSE) for each recorded field location.")

set_heading(doc, "Step 1 — Spatial and Temporal Grouping", level=2)
add_para(doc,
    "The 1,254 field rows were grouped into spatial cells of approximately 2 km x 2 km. "
    "Rows falling within the same spatial cell and captured within a 14-day window "
    "were assigned to a single Sentinel-2 tile. This grouping strategy reduced the "
    "number of required tile downloads from 1,254 to just 8, significantly minimizing "
    "storage and download time.")

set_heading(doc, "Step 2 — Product Search via Copernicus Catalog", level=2)
add_para(doc,
    "For each spatial-temporal group, the researchers queried the Copernicus OData "
    "catalog to identify the most suitable Sentinel-2 Level-2A (L2A) product. "
    "L2A products provide bottom-of-atmosphere (BOA) reflectance, meaning atmospheric "
    "effects have already been corrected. The search criteria prioritized cloud-free "
    "scenes closest in acquisition date to the field capture date.")
add_code(doc, "Catalog API: https://catalogue.dataspace.copernicus.eu/odata/v1/Products")

set_heading(doc, "Step 3 — Product Download via CDSE S3", level=2)
add_para(doc,
    "Each identified product (packaged as a .SAFE directory, approximately 908 MB per tile) "
    "was downloaded using CDSE's S3-compatible storage endpoint. The S3 method was selected "
    "over the standard HTTP zipper endpoint because it allows per-file resume capability, "
    "avoiding full re-downloads in the event of connection interruptions.")
add_code(doc, "S3 Endpoint: https://eodata.dataspace.copernicus.eu  (Bucket: eodata)")
add_para(doc,
    "Authentication was performed using the OAuth2 password grant flow with "
    "client_id=cdse-public, using the researcher's registered Copernicus "
    "username and password credentials stored in an environment file (.env).",
    indent=True)

set_heading(doc, "Step 4 — Spectral Band Sampling at Field Points", level=2)
add_para(doc,
    "Once a tile was downloaded, the pipeline opened each spectral band file "
    "(.jp2 format) using the rasterio geospatial library. For each field point, "
    "the GPS coordinates (WGS84 latitude/longitude) were reprojected into the tile's "
    "native UTM coordinate system. A 3x3 pixel window centered on the reprojected "
    "point was then sampled and averaged to reduce single-pixel noise. "
    "This process was repeated for all 12 Sentinel-2 bands.")

add_divider(doc)
add_table(doc,
    ["Band", "Wavelength", "Spatial Resolution", "Spectral Domain"],
    [
        ("B01", "443 nm",  "60 m", "Coastal aerosol"),
        ("B02", "490 nm",  "10 m", "Blue"),
        ("B03", "560 nm",  "10 m", "Green"),
        ("B04", "665 nm",  "10 m", "Red"),
        ("B05", "705 nm",  "20 m", "Red Edge 1 — sensitive to chlorophyll and phosphorus"),
        ("B06", "740 nm",  "20 m", "Red Edge 2"),
        ("B07", "783 nm",  "20 m", "Red Edge 3 — nitrogen-sensitive"),
        ("B08", "842 nm",  "10 m", "Near-Infrared (NIR) — vegetation density"),
        ("B8A", "865 nm",  "20 m", "Narrow NIR — chlorophyll"),
        ("B09", "945 nm",  "60 m", "Water vapour"),
        ("B11", "1610 nm", "20 m", "Short-Wave Infrared 1 — soil moisture and potassium"),
        ("B12", "2190 nm", "20 m", "Short-Wave Infrared 2"),
    ]
)
add_divider(doc)

set_heading(doc, "Step 5 — Incremental Output", level=2)
add_para(doc,
    "Following the successful processing of each tile, the corresponding field rows "
    "with appended band values were immediately written to the output CSV file. "
    "This incremental approach ensured that partial results were preserved in the "
    "event of a processing interruption. Download progress was monitored in real "
    "time via a terminal progress bar (tqdm).")
add_code(doc, "data/processed/field_data_with_bands.csv")
add_para(doc,
    "Current status: 410 rows have been saved following the completion of the first "
    "of eight scheduled tile downloads.",
    indent=True)

add_divider(doc)

# ── Phase 3 ───────────────────────────────────────────────────────────────────
set_heading(doc, "Phase 3 — Soil Ground Truth Collection via Rapid Test Kit", level=1)

add_para(doc,
    "In parallel with satellite data acquisition, the researchers collected soil ground "
    "truth measurements at each field spot using a commercial Rapid Soil Test Kit (STK). "
    "The kit produces color-based reagent readings that are matched against reference "
    "color charts to determine the nutrient and pH level of each soil sample.")

add_divider(doc)
add_table(doc,
    ["Soil Property", "Measurement Type", "Ground Truth Categories"],
    [
        ("Nitrogen (N)",   "Ordinal 3-class",   "Low / Medium / High"),
        ("Phosphorus (P)", "Ordinal 3-class",   "Low / Medium / High"),
        ("Potassium (K)",  "Ordinal 3-class",   "Low / Medium / High"),
        ("Soil pH",        "11-class CPR scale",
         "4.0, 4.4, 4.8, 5.2, 5.4, 5.8, 6.0, 6.4, 6.8, 7.2, 7.6"),
    ]
)
add_divider(doc)

add_para(doc,
    "The pH values are read from the Color Plate Reader (CPR) strip included in the "
    "rapid test kit, which distinguishes 11 discrete pH levels calibrated to the "
    "acidic to near-neutral range typical of highland agricultural soils in Benguet.")

set_heading(doc, "Required Data Format for Merging", level=2)
add_para(doc,
    "The STK results are recorded with the unique identifier (uuid) from the "
    "AgriCapture field data, enabling a direct join between ground truth and "
    "the extracted satellite band values:")
add_code(doc, "uuid, nitrogen, phosphorus, potassium, ph")
add_code(doc, "aeeed204, Medium, Low, High, 5.8")
add_code(doc, "a18676dc, Low, Medium, Medium, 6.0")

add_para(doc,
    "Once the STK results are available in this format, the merge script "
    "produces the final training-ready dataset by joining on the uuid column:",
    indent=True)
add_code(doc, "python src/merge_stk.py  stk_results.csv")

add_divider(doc)

# ── Phase 4 ───────────────────────────────────────────────────────────────────
set_heading(doc, "Phase 4 — Feature Engineering (Spectral Indices)", level=1)

add_para(doc,
    "Prior to model training, the researchers augmented the raw spectral band values "
    "with ten computed spectral indices. These derived features are established in "
    "the remote sensing literature as proxies for soil and vegetation conditions "
    "that raw reflectance values do not capture independently.")

add_divider(doc)
add_table(doc,
    ["Index", "Formula", "Soil/Vegetation Relevance"],
    [
        ("NDVI",   "(B08 - B04) / (B08 + B04)",
         "Normalized Difference Vegetation Index — canopy density and biomass"),
        ("EVI",    "2.5 x (B08 - B04) / (B08 + 6*B04 - 7.5*B02 + 1)",
         "Enhanced Vegetation Index — aerosol-corrected canopy signal"),
        ("SAVI",   "1.5 x (B08 - B04) / (B08 + B04 + 0.5)",
         "Soil-Adjusted Vegetation Index — minimizes bare soil influence"),
        ("MSAVI",  "Modified SAVI (variable soil factor)",
         "Improved on sparse or partly bare canopy"),
        ("NDRE",   "(B8A - B05) / (B8A + B05)",
         "Normalized Difference Red Edge — chlorophyll and nitrogen stress"),
        ("CHL_re", "(B8A / B05) - 1",
         "Chlorophyll Red-Edge Index — crop health indicator"),
        ("BSI",    "(B11+B04 - B08-B02) / (B11+B04 + B08+B02)",
         "Bare Soil Index — fraction of exposed soil"),
        ("BI",     "sqrt((B04^2 + B08^2) / 2)",
         "Brightness Index — soil surface reflectance"),
        ("NDWI",   "(B03 - B08) / (B03 + B08)",
         "Normalized Difference Water Index — surface water content"),
        ("NDMI",   "(B08 - B11) / (B08 + B11)",
         "Normalized Difference Moisture Index — canopy/soil moisture"),
    ]
)
add_divider(doc)

set_heading(doc, "Final Feature Vector", level=2)
add_para(doc,
    "Each training sample was represented by the following 26-dimensional feature vector:")
add_bullet(doc, "12 raw Sentinel-2 spectral bands (B01 through B12 including B8A)")
add_bullet(doc, "10 computed spectral indices (as listed above)")
add_bullet(doc, "3 microclimate features: temperature (C), humidity (%), altitude (m)")
add_bullet(doc, "1 categorical feature: crop type (one-hot encoded at preprocessing)")

add_divider(doc)

# ── Phase 5 ───────────────────────────────────────────────────────────────────
set_heading(doc, "Phase 5 — Machine Learning Model Training", level=1)

add_para(doc,
    "The researchers trained and compared three machine learning classifiers on the "
    "merged field-satellite-STK dataset. Each model independently predicted all four "
    "soil properties — Nitrogen, Phosphorus, Potassium, and pH. "
    "The comparative evaluation was conducted to identify the best-performing model "
    "for each soil property and to provide a multi-model analysis suitable for "
    "inclusion in a research paper.")

set_heading(doc, "Models Evaluated", level=2)
add_table(doc,
    ["Model", "Configuration", "Rationale"],
    [
        ("XGBoost",
         "500 estimators, learning rate 0.03, max depth 6, L1/L2 regularisation",
         "Strong performance on tabular data; handles feature interactions well"),
        ("Random Forest",
         "500 trees, sqrt max features, balanced class weights",
         "Ensemble robustness; effective on small datasets with noise"),
        ("SVM (RBF kernel)",
         "C=10, gamma=scale, one-vs-rest multiclass strategy",
         "Well-suited for high-dimensional scaled feature spaces"),
    ]
)

set_heading(doc, "Spatial Cross-Validation", level=2)
add_para(doc,
    "To prevent data leakage arising from spatially correlated GPS points, the "
    "researchers applied GroupKFold cross-validation with barangay as the grouping "
    "variable. This ensured that field spots from the same barangay were never "
    "present in both the training and test sets simultaneously, producing a more "
    "realistic estimate of generalization performance across unseen farm locations. "
    "Five-fold cross-validation was used, reduced to the number of unique barangays "
    "when fewer than five were present in the dataset.")

set_heading(doc, "Class Imbalance Handling", level=2)
add_para(doc,
    "Due to the unequal distribution of soil nutrient classes (e.g., a higher "
    "proportion of Medium nitrogen samples than Low or High), the researchers applied "
    "balanced sample weighting during XGBoost and Random Forest training. "
    "Weights were computed proportionally to inverse class frequency using "
    "sklearn's compute_sample_weight function, ensuring minority classes "
    "contributed fairly to the loss function.")

set_heading(doc, "Evaluation Metrics", level=2)
add_para(doc,
    "The following metrics were computed for each model and soil property target, "
    "reported as per-fold mean +/- standard deviation and pooled across all folds:")
add_divider(doc)
add_table(doc,
    ["Metric", "Description and Research Relevance"],
    [
        ("Overall Accuracy (OA)",
         "Percentage of correctly classified samples across all classes"),
        ("Macro F1-Score",
         "Average F1 across all classes with equal weight — accounts for class imbalance"),
        ("Weighted F1-Score",
         "F1 weighted by class frequency — reflects overall dataset performance"),
        ("Cohen's Kappa",
         "Agreement above chance; <0.40 = fair, 0.40-0.60 = moderate, 0.60+ = substantial"),
        ("Ordinal MAE",
         "Mean absolute error in class steps — captures ordinal prediction quality"),
        ("pH MAE (pH units)",
         "Ordinal MAE converted to real pH scale units (each step = 0.4 pH units)"),
        ("Per-fold Mean +/- Std",
         "Demonstrates result consistency and stability across spatial folds"),
        ("95% Confidence Interval",
         "Statistical reliability bound computed via t-distribution for publication"),
    ]
)
add_divider(doc)

set_heading(doc, "Run Command", level=3)
add_code(doc, "# Training on synthetic data (pipeline testing):")
add_code(doc, "python src/train_ordinal.py data/processed/synthetic_test.csv")
add_code(doc, "")
add_code(doc, "# Training on real field data (requires STK ground truth):")
add_code(doc, "python src/train_ordinal.py data/processed/field_data_with_bands.csv "
              "--figures-dir outputs/real/figures --output-dir outputs/real")

add_divider(doc)

# ── Phase 6 ───────────────────────────────────────────────────────────────────
set_heading(doc, "Phase 6 — Output Files and Results", level=1)

add_para(doc,
    "Upon completion of training, the pipeline automatically saves all results "
    "to the designated output directory. The outputs are structured to support "
    "direct inclusion in a research paper.")

add_table(doc,
    ["Output File", "Description"],
    [
        ("outputs/real/metrics_summary.csv",
         "Consolidated metrics table for all models across all soil property targets. "
         "One row per model-target combination. Suitable for direct use as a results table."),
        ("outputs/real/figures/model_comparison_<target>.png",
         "Grouped bar chart comparing Overall Accuracy, Macro F1, and Cohen's Kappa "
         "across XGBoost, Random Forest, and SVM for a given target."),
        ("outputs/real/figures/confusion_matrix_<target>_<model>.png",
         "Row-normalized confusion matrix per model. 11x11 grid for pH (CPR scale); "
         "3x3 grid for N, P, and K. Color intensity represents per-class recall."),
        ("outputs/real/figures/feature_importance_<target>_<model>.png",
         "Horizontal bar chart of the top 20 most important features for the "
         "best-performing tree-based model for each target."),
    ]
)

add_divider(doc)

# ── Status Summary ─────────────────────────────────────────────────────────────
set_heading(doc, "Current Pipeline Status", level=1)

add_table(doc,
    ["Pipeline Step", "Status", "Notes"],
    [
        ("Field data collection (AgriCapture)", "Completed",
         "1,254 rows recorded across 4 barangays in Benguet"),
        ("Sentinel-2 tile download — Tile 1 of 8", "Completed",
         "908 MB downloaded via CDSE S3 in approximately 1 hour 40 minutes"),
        ("Sentinel-2 tile download — Tiles 2 to 8", "In Progress",
         "Approximately 7 x 908 MB remaining"),
        ("Spectral band sampling — Tile 1", "Completed",
         "410 field rows augmented with 12 band values"),
        ("STK ground truth collection", "Required",
         "Soil test kit readings must be recorded and formatted as CSV"),
        ("STK and satellite data merge", "Blocked",
         "Dependent on STK CSV file with uuid column"),
        ("Model training on real data", "Blocked",
         "Dependent on merged dataset containing STK labels"),
        ("Model training on synthetic data", "Operational",
         "Pipeline tested and validated using generated synthetic dataset"),
    ]
)

# ── Save ──────────────────────────────────────────────────────────────────────
os.makedirs(os.path.dirname(OUTPUT_PATH), exist_ok=True)
doc.save(OUTPUT_PATH)
print(f"Saved: {OUTPUT_PATH}")
